"""Generates the full assignment report (submission/Report.docx),
structured to mirror the brief's own task numbering (Tasks 1-6 + Bonus
+ Required Deliverables) so a marker can map sections to the marking
scheme directly. Not part of the ontology/backend/frontend deliverable
itself - just tooling, kept for reproducibility."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()


def set_cell_shading(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def add_table(headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(hdr[i], "1F4E78")
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    return table


def code_para(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.25)
    return p


# ============================== HEADER ==============================
doc.add_heading("Local-Domain Ontology Design, Reasoning and SPARQL", level=0)
meta = doc.add_paragraph()
for label, value in [
    ("Index No: ", "258843A"),
    ("Name: ", "Weerarathna T.M.P."),
    ("Module: ", "IT5432 - Semantic Web and Ontological Engineering"),
    ("Domain: ", "Abans - Sri Lankan electronics and home appliance retail (buyabans.com)"),
]:
    meta.add_run(label).bold = True
    meta.add_run(value + "\n")
meta.add_run("GitHub Repository: ").bold = True
gh = meta.add_run("https://github.com/PathmikaW/protege_exercise/tree/main/assignment")
gh.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
gh.font.underline = True

doc.add_paragraph()

# ============================== TASK 1 ==============================
doc.add_heading("Task 1 - Domain Selection and Requirements Analysis", level=1)

doc.add_heading("Domain Description", level=2)
doc.add_paragraph(
    "The ontology models the real product catalog of Abans (buyabans.com), one of "
    "Sri Lanka's largest electronics and home-appliance retailers (40+ brands, "
    "400+ showrooms island-wide). It covers four of Abans' actual top-level product "
    "categories - Television, Home Appliances, Mobile Devices, and Computers - along "
    "with Abans' real multi-branch showroom network (city and province locations), "
    "which is a genuinely local, industry-relevant domain grounded in verifiable, "
    "real-world data rather than an invented scenario."
)

doc.add_heading("Purpose and Expected Users", level=2)
doc.add_paragraph(
    "The ontology is designed to support product discovery and comparison for a "
    "retail context: helping a customer or store assistant find products by brand, "
    "price, feature set, energy efficiency, and showroom availability, and to "
    "automatically classify products into useful marketing/business categories "
    "(e.g. Smart, Premium, Eco-Friendly) without those classifications needing to "
    "be manually tagged on every product."
)

doc.add_heading("Scope and Boundaries", level=2)
doc.add_paragraph(
    "In scope: the four product categories listed above, their real sub-types, "
    "brands, key attributes (price, warranty, capacity, screen size, features, "
    "energy rating), and the showroom/location network. Out of scope: the "
    "remaining nine Abans catalog categories (Kitchen Appliances, Fashion, "
    "MINISO, etc.), customer accounts, ordering/checkout, and stock-level "
    "quantities - these were excluded to keep the ontology focused and "
    "manageable while still being rich enough to demonstrate every required "
    "OWL construct meaningfully."
)

doc.add_heading("Competency Questions", level=2)
doc.add_paragraph("At least five were required; fifteen were developed and implemented as SPARQL (see Task 5):")
cqs = [
    "What products does a given brand sell?",
    "What TVs cost less than Rs. 150,000?",
    "Which products have WiFi?",
    "Which products have 2+ years warranty?",
    "Which products are classified as Smart?",
    "Which products are Eco-Friendly (A++/A+++)?",
    "Which products are Premium (price >= 150,000)?",
    "Which products are Budget (complement of Premium)?",
    "How many products exist per category?",
    "What is the average price per category?",
    "Which brands appear in more than one category?",
    "Which products have 3+ features (Fully-Loaded)?",
    "Which products are available at showrooms in the Western Province?",
    "Which showrooms stock a given product, and in which city/province?",
    "List all products with their capacity, where available.",
]
for c in cqs:
    doc.add_paragraph(c, style="List Number")

# ============================== TASK 2 ==============================
doc.add_heading("Task 2 - Ontology Design", level=1)
doc.add_paragraph(
    "The final ontology contains 47 classes, 15 properties (10 object + 5 data), "
    "59 individuals, and 598 axioms (461 logical), verified consistent by the "
    "HermiT reasoner."
)

doc.add_heading("Main Classes and Class Hierarchy", level=2)
doc.add_paragraph(
    "Product is the root class, with four mutually disjoint sub-categories - TV, "
    "HomeAppliance, MobileDevice, Computer - each further divided into disjoint "
    "real sub-types (e.g. TV -> LEDTV, SmartLEDTV, UHDTV, OLEDTV, QLEDTV). "
    "Supporting classes include Brand, Feature (with disjoint sub-types "
    "SmartFeature/WiFiFeature/InverterFeature/BluetoothFeature), "
    "EnergyRatingPartition (a five-level value partition), and a Location "
    "hierarchy (Province, City, Showroom) modeling Abans' real branch network."
)

doc.add_heading("Object and Data Properties", level=2)
add_table(
    ["Property", "Domain", "Range", "Characteristics"],
    [
        ("hasBrand", "Product", "Brand", "Functional"),
        ("hasFeature", "Product", "Feature", "-"),
        ("hasEnergyRating", "Product", "EnergyRatingPartition", "Functional"),
        ("locatedIn", "Location", "Location", "Transitive"),
        ("contains", "Location", "Location", "Transitive (inverse of locatedIn)"),
        ("availableAt", "Product", "Showroom", "-"),
        ("stocks", "Showroom", "Product", "Inverse of availableAt"),
        ("hasPrice", "Product", "xsd:decimal", "Functional"),
        ("hasWarrantyYears", "Product", "xsd:integer", "Functional"),
        ("hasCapacity", "Product", "xsd:decimal", "-"),
    ],
)

doc.add_heading("Representative Individuals", level=2)
doc.add_paragraph(
    "18 real product individuals (e.g. iPhone_15_Pro, Samsung_Inverter_AC_12000BTU, "
    "LG_65_OLED_TV), 10 Brand individuals (LG, Samsung, Apple, Whirlpool, Haier, HP, "
    "Xiaomi, Lenovo, Toshiba, Philips), and 9 location individuals (3 Provinces, "
    "3 Cities, 3 Showrooms) modeling Abans' real branch presence in Colombo, "
    "Kandy, and Galle."
)

doc.add_heading("Important Restrictions", level=2)
doc.add_paragraph(
    "Every restriction type required by the brief is used, each justified by a "
    "genuine business rule rather than added to satisfy a count:"
)
restr = [
    ("Necessary condition", "Product ⊑ hasBrand some Brand", "every product must have a brand"),
    ("Existential restriction", "SmartProduct ≡ Product ⊓ hasFeature some SmartFeature", "at least one smart feature"),
    ("Universal restriction + closure", "FullySmartProduct ≡ Product ⊓ hasFeature only (SmartFeature ⊔ WiFiFeature ⊔ BluetoothFeature)", "every feature must be smart-related"),
    ("Datatype facet restriction", "PremiumProduct ≡ Product ⊓ hasPrice some xsd:decimal[>=150000]", "numeric price threshold"),
    ("Complement class", "BudgetProduct ≡ Product ⊓ ¬PremiumProduct", "not premium"),
    ("Cardinality restriction", "FullyLoadedProduct ≡ Product ⊓ hasFeature min 3", "3 or more distinct features"),
    ("Nested restriction", "ColomboAvailableProduct ≡ Product ⊓ availableAt some (Showroom ⊓ locatedIn value Colombo)", "available at a Colombo showroom"),
]
add_table(["Restriction type", "Example axiom", "Business justification"], restr)

doc.add_heading("Ontology Diagram", level=2)
doc.add_paragraph(
    "The diagram below shows the principal concepts and relationships: the "
    "disjoint category hierarchy, the property links to Brand/Feature/"
    "EnergyRatingPartition, the transitive showroom-location network, and the "
    "reasoner-defined classes (dashed) that individuals are classified into "
    "automatically."
)
doc.add_picture("submission/ontology_diagram.png", width=Inches(6.5))

# ============================== TASK 3 ==============================
doc.add_heading("Task 3 - Ontology Development", level=1)
doc.add_paragraph(
    "The ontology was implemented in Protégé 5.6.9. Every class has a "
    "clean PascalCase identifier plus an rdfs:label carrying Abans' exact real "
    "site wording (e.g. class HomeAppliance, label \"Home Appliances\"), so the "
    "ontology is technically well-formed while displaying authentic real-world "
    "terminology. Major classes carry rdfs:comment annotations documenting their "
    "purpose. Domains and ranges are specified on every object and data property "
    "(see Task 2 table)."
)
doc.add_paragraph(
    "OWL features used: equivalent classes (9 defined classes), disjoint classes "
    "(11 groups), inverse properties (3 pairs), transitive properties (locatedIn/"
    "contains), functional properties (hasBrand, hasEnergyRating, hasPrice, "
    "hasWarrantyYears), existential and universal restrictions, a datatype facet "
    "restriction, a cardinality restriction, and a hasValue restriction "
    "(PremiumAppleProduct's hasBrand value Apple). No feature was added purely to "
    "satisfy a count - each is tied to a genuine Abans business concept, as shown "
    "in the Task 2 justification table."
)

# ============================== TASK 4 ==============================
doc.add_heading("Task 4 - Reasoning and Inference", level=1)
doc.add_paragraph(
    "The HermiT reasoner (Protege's built-in reasoner) was run against the "
    "completed ontology. Result: the ontology is logically consistent, with no "
    "unsatisfiable classes."
)
doc.add_heading("At Least Five Inferred Facts", level=2)
doc.add_paragraph(
    "None of the facts below were directly asserted when the individuals were "
    "created - each is a consequence the reasoner derived from the axioms."
)
inferred = [
    ("iPhone_15_Pro is a SmartProduct", "SmartProduct ≡ Product ⊓ hasFeature some SmartFeature; iPhone_15_Pro has a SmartFeature asserted"),
    ("iPhone_15_Pro is a PremiumProduct", "PremiumProduct's datatype facet hasPrice some xsd:decimal[>=150000]; iPhone_15_Pro's price is 385,000"),
    ("iPhone_15_Pro is a FullyLoadedProduct", "Cardinality restriction hasFeature min 3; iPhone_15_Pro has exactly 3 features asserted"),
    ("LG_65_OLED_TV is an EcoFriendlyProduct", "hasEnergyRating some (A++ or A+++); LG_65_OLED_TV's rating is A++"),
    ("Samsung_Inverter_AC_12000BTU is an InverterHomeAppliance", "Combined restriction requiring an InverterFeature and a high energy rating, both asserted on this individual"),
    ("Abans_Colombo_City is locatedIn WesternProvince", "locatedIn is Transitive; only Showroom->City (Colombo) and City->Province (WesternProvince) were asserted directly - the Showroom->Province link is inferred"),
    ("Mystery_Smart_Gadget is a SmartProduct but NOT a FullySmartProduct", "It has one SmartFeature asserted (satisfies the existential), but no closure axiom rules out other undisclosed features, so the reasoner correctly refuses to classify it as fully-smart (Open World Assumption)"),
]
add_table(["Inferred fact", "Axiom(s) that produced it"], inferred)
doc.add_paragraph(
    "Evidence: screenshots of the Protégé Individuals (Inferred) view for each of "
    "the above are attached at the end of this report."
)

# ============================== TASK 5 ==============================
doc.add_heading("Task 5 - SPARQL Query Development", level=1)
doc.add_paragraph(
    "Fifteen SPARQL queries were developed (minimum required: eight), covering "
    "every required category: two basic SELECT queries (CQ1, CQ3), a FILTER "
    "query (CQ2), an OPTIONAL query (CQ15), an aggregation query (CQ9/CQ10 - "
    "COUNT and AVG), grouping (CQ9/CQ10 use GROUP BY), and multiple queries that "
    "retrieve reasoner-derived knowledge (CQ5-CQ8, CQ12 via defined-class "
    "reasoning; CQ13-CQ14 via the transitive locatedIn property)."
)
doc.add_paragraph(
    "Technical note on asserted vs. inferred: reasoner-derived facts (e.g. class "
    "membership in SmartProduct, or the Showroom->Province transitive link) were "
    "not queryable by a plain SPARQL engine, which has no reasoner of its own. "
    "These were materialised using Protégé's 'Export inferred axioms as ontology' "
    "feature, which physically writes the reasoner's conclusions into the "
    "exported OWL file as real triples - the approach the brief's technical note "
    "explicitly recommends."
)

query_data = [
    ("CQ1", "What products does brand X sell?",
     "SELECT ?product WHERE {\n  ?product abans:hasBrand abans:LG .\n}",
     "LG_43_LED_TV, LG_65_OLED_TV", "Asserted",
     "Two LG products found via the directly asserted hasBrand relationship."),
    ("CQ2", "What TVs cost less than Rs. 150,000?",
     "SELECT ?product ?price WHERE {\n  ?product a abans:TV .\n  ?product abans:hasPrice ?price .\n  FILTER(?price < 150000)\n}",
     "LG_43_LED_TV (65000), Samsung_55_Smart_LED_TV (145000), Toshiba_50_UHD_TV (98000)", "Asserted",
     "Three of the four TVs qualify; LG_65_OLED_TV (350000) is correctly excluded."),
    ("CQ3", "Which products have WiFi?",
     "SELECT ?product WHERE {\n  ?product abans:hasFeature ?feature .\n  ?feature a abans:WiFiFeature .\n}",
     "7 products (LG_65_OLED_TV, Toshiba_50_UHD_TV, iPhone_15_Pro, HP_Pavilion_Laptop, Lenovo_IdeaPad_Laptop, Lenovo_Desktop_Monitor_Combo, HP_Tablet)", "Asserted",
     "Simple join across all products with an asserted WiFiFeature."),
    ("CQ4", "Which products have 2+ years warranty?",
     "SELECT ?product ?years WHERE {\n  ?product abans:hasWarrantyYears ?years .\n  FILTER(?years >= 2)\n}",
     "8 products, warranty 2-5 years", "Asserted",
     "Filter on the directly asserted hasWarrantyYears value."),
    ("CQ5", "Which products are classified as Smart?",
     "SELECT ?product WHERE {\n  ?product a abans:SmartProduct .\n}",
     "7 products including Mystery_Smart_Gadget", "Inferred",
     "None of these were typed SmartProduct directly - membership comes from the equivalentClass axiom."),
    ("CQ6", "Which products are Eco-Friendly (A++/A+++)?",
     "SELECT ?product WHERE {\n  ?product a abans:EcoFriendlyProduct .\n}",
     "LG_65_OLED_TV, Samsung_Inverter_AC_12000BTU, Whirlpool_Front_Load_Washing_Machine", "Inferred",
     "Classified via the energy-rating existential restriction."),
    ("CQ7", "Which products are Premium (price >= 150,000)?",
     "SELECT ?product ?price WHERE {\n  ?product a abans:PremiumProduct .\n  ?product abans:hasPrice ?price .\n}",
     "6 products, 165,000-385,000", "Inferred",
     "Classified via the datatype facet restriction on price."),
    ("CQ8", "Which products are Budget (complement of Premium)?",
     "SELECT ?product WHERE {\n  ?product a abans:BudgetProduct .\n}",
     "12 products", "Inferred",
     "Complement-class reasoning: everything that is provably not Premium."),
    ("CQ9", "How many products exist per category?",
     "SELECT ?category (COUNT(?product) AS ?count) WHERE {\n  ?product a ?category .\n  FILTER(?category IN (abans:TV, abans:HomeAppliance, abans:MobileDevice, abans:Computer))\n}\nGROUP BY ?category",
     "TV: 4, HomeAppliance: 4, MobileDevice: 6, Computer: 4", "Asserted",
     "Aggregate (COUNT) with GROUP BY over asserted category types."),
    ("CQ10", "What is the average price per category?",
     "SELECT ?category (AVG(?price) AS ?avgPrice) WHERE {\n  ?product a ?category .\n  ?product abans:hasPrice ?price .\n  FILTER(?category IN (abans:TV, abans:HomeAppliance, abans:MobileDevice, abans:Computer))\n}\nGROUP BY ?category",
     "Computer: 120000, HomeAppliance: 159750, TV: 164500, MobileDevice: 104250", "Asserted",
     "Aggregate (AVG) with GROUP BY."),
    ("CQ11", "Which brands appear in more than one category?",
     "SELECT ?brand (COUNT(DISTINCT ?category) AS ?catCount) WHERE {\n  ?product abans:hasBrand ?brand .\n  ?product a ?category .\n  FILTER(?category IN (abans:TV, abans:HomeAppliance, abans:MobileDevice, abans:Computer))\n}\nGROUP BY ?brand\nHAVING (COUNT(DISTINCT ?category) > 1)",
     "Samsung (3 categories)", "Asserted",
     "Multi-hop join with HAVING; Samsung genuinely spans TV, HomeAppliance and MobileDevice."),
    ("CQ12", "Which products have 3+ features (Fully-Loaded)?",
     "SELECT ?product WHERE {\n  ?product a abans:FullyLoadedProduct .\n}",
     "LG_65_OLED_TV, iPhone_15_Pro", "Inferred",
     "Classified via the cardinality restriction (hasFeature min 3)."),
    ("CQ13", "Which products are available in the Western Province?",
     "SELECT ?product ?showroom WHERE {\n  ?product abans:availableAt ?showroom .\n  ?showroom abans:locatedIn abans:WesternProvince .\n}",
     "10 products, all via Abans_Colombo_City", "Inferred",
     "Works as a single hop because locatedIn's transitive closure was materialised."),
    ("CQ14", "Which showrooms stock a given product, in which city/province?",
     "SELECT ?showroom ?city ?province WHERE {\n  abans:LG_43_LED_TV abans:availableAt ?showroom .\n  ?showroom abans:locatedIn ?city .\n  ?city a abans:City .\n  ?city abans:locatedIn ?province .\n  ?province a abans:Province .\n}",
     "Abans_Colombo_City, Colombo, WesternProvince", "Inferred",
     "Multi-hop join over the transitive locatedIn property."),
    ("CQ15", "List all products with their capacity, where available",
     "SELECT ?product ?capacity WHERE {\n  ?product a abans:Product .\n  OPTIONAL { ?product abans:hasCapacity ?capacity }\n}",
     "18 rows; capacity bound for 4 appliances, unbound elsewhere", "Asserted",
     "OPTIONAL correctly returns every product even when hasCapacity doesn't apply."),
]

for cq_id, question, sparql, result, kind, interp in query_data:
    doc.add_heading(f"{cq_id} - {question}", level=3)
    code_para(sparql)
    p = doc.add_paragraph()
    p.add_run("Result: ").bold = True
    p.add_run(result + "\n")
    p.add_run("Asserted or Inferred: ").bold = True
    r = p.add_run(kind)
    r.font.color.rgb = RGBColor(0xC6, 0x28, 0x28) if kind == "Inferred" else RGBColor(0x2E, 0x7D, 0x32)
    r.bold = True
    p.add_run("\n")
    p.add_run("Interpretation: ").bold = True
    p.add_run(interp)

# ============================== TASK 6 ==============================
doc.add_heading("Task 6 - Evaluation and Documentation", level=1)
eval_points = [
    ("Logical consistency", "Verified with HermiT after every major change; the ontology is fully consistent with no unsatisfiable classes."),
    ("Correctness of class hierarchy", "Every sibling group (categories, sub-types, features, energy ratings) is explicitly disjoint, preventing accidental overlap; the hierarchy was cross-checked against Abans' real site navigation."),
    ("Coverage of competency questions", "All 15 competency questions were implemented and verified to return correct results (Task 5)."),
    ("Quality of naming and documentation", "Clean PascalCase class/property identifiers with rdfs:label annotations carrying real Abans wording, plus rdfs:comment on major classes for readability."),
    ("Reusability and extensibility", "The category/sub-type pattern and the defined-class pattern (Smart/Premium/Eco-Friendly/etc.) can be extended to Abans' remaining nine catalog categories, or adapted to a similar retailer, with minimal restructuring."),
    ("Limitations", "Scope is limited to 4 of Abans' 13 real categories; prices are modeled as plain decimals without currency/locale typing; stock quantities and multi-currency pricing are not modeled; the ontology does not model customer accounts or transactions."),
    ("Future improvements", "Extend to the remaining catalog categories; add a Currency/Price value object; add temporal properties for stock/discount validity; integrate live product data via a scheduled import instead of static individuals."),
]
for title, text in eval_points:
    p = doc.add_paragraph()
    p.add_run(title + ": ").bold = True
    p.add_run(text)

# ============================== BONUS ==============================
doc.add_heading("Bonus Task - Application Integration", level=1)
doc.add_paragraph(
    "A full-stack application was built on top of the ontology: a FastAPI + "
    "rdflib backend, and a Streamlit front end."
)
doc.add_heading("Minimum Capability Coverage", level=2)
doc.add_paragraph(
    "The application covers more than the required minimum of two capabilities:"
)
for cap in [
    "Run predefined SPARQL queries - all 15 competency questions, selectable from a dropdown.",
    "Display inferred classifications - CQ5-CQ8, CQ12 all surface reasoner-derived class membership.",
    "Enter criteria and retrieve matching knowledge - the Raw SPARQL tab accepts any ad-hoc query.",
]:
    doc.add_paragraph(cap, style="List Bullet")

doc.add_heading("Architecture", level=2)
doc.add_paragraph(
    "backend/main.py loads the exported abans.owl into an in-memory rdflib "
    "graph at startup and exposes REST endpoints: GET /health, GET "
    "/competency-questions, GET /competency-questions/{id}/run, and POST "
    "/sparql for raw queries. frontend/app.py (Streamlit) calls these endpoints "
    "over HTTP and renders results as a table, with a live connection-status "
    "and ontology-statistics sidebar."
)

doc.add_heading("Setup and Execution Instructions", level=2)
code_para(
    "cd assignment\n"
    "python -m venv .venv\n"
    ".venv\\Scripts\\Activate.ps1\n"
    "pip install -r requirements.txt\n\n"
    ".venv\\Scripts\\python.exe -m uvicorn backend.main:app --port 8010\n"
    ".venv\\Scripts\\python.exe -m streamlit run frontend/app.py --server.port 8511"
)

doc.add_heading("Evidence of Reasoning-Dependent Results", level=2)
doc.add_paragraph(
    "Running CQ5 (SmartProduct), CQ6 (EcoFriendlyProduct), CQ7 (PremiumProduct), "
    "CQ8 (BudgetProduct) or CQ12 (FullyLoadedProduct) through the application "
    "returns results that depend entirely on reasoner-inferred class membership, "
    "materialised into the exported OWL file as described in Task 5. Screenshots "
    "of these queries running in the application are attached at the end of "
    "this report."
)

# ============================== CONCLUSION / REFS ==============================
doc.add_heading("Conclusion", level=1)
doc.add_paragraph(
    "This assignment models a real, verifiable local-industry domain (Abans' "
    "electronics retail catalog) as an OWL ontology, demonstrating the full "
    "range of required modeling constructs, reasoner-driven inference including "
    "the Open World Assumption, and knowledge retrieval through fifteen SPARQL "
    "competency questions. A working full-stack application was additionally "
    "built to query the ontology through a web interface, going beyond the "
    "assignment's bonus minimum requirement."
)

doc.add_heading("References", level=1)
refs = [
    "Abans / buyabans.com - product category, brand, and attribute data (accessed 2026-08-28).",
    "Protege Desktop 5.6.9 - ontology editor and HermiT reasoner.",
    "rdflib (Python) - RDF graph and SPARQL query engine.",
    "FastAPI - backend web framework.",
    "Streamlit - front-end application framework.",
]
for r in refs:
    doc.add_paragraph(r, style="List Bullet")

doc.add_heading("AI-Assistance Acknowledgment", level=1)
doc.add_paragraph(
    "Portions of this work were developed with the assistance of Claude "
    "(Anthropic), used as a modelling and coding aid throughout the process - "
    "similarly to how one might use a tutorial, a pair-programming partner, or "
    "reference documentation. All domain analysis choices, the selection of "
    "Abans as the domain, the specific classes/properties/restrictions "
    "modelled, and the interpretation of every inference and query result are "
    "my own, and I am able to explain and demonstrate every part of the "
    "submitted work as required by the assignment's General Instructions."
)

doc.add_heading("Screenshots", level=1)
doc.add_paragraph("(Screenshots for Task 4 evidence and the Bonus application are attached below.)")

doc.save("submission/258843A_Assignment_Report.docx")
print("Report generated.")
